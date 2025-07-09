import pandas as pd
import requests
import logging
from flask import current_app
from typing import Dict, List, Any, Optional, Tuple
import io
import json
import re
from datetime import datetime

# PDF processing
import PyPDF2
import pdfplumber

# DOCX processing
from docx import Document

# Advanced CSV processing
import csv

class EnhancedFileProcessor:
    """Enhanced file processor supporting CSV, PDF, DOCX with intelligent content extraction"""
    
    def __init__(self, db_manager):
        self.db_manager = db_manager
        self.supported_formats = ['csv', 'xlsx', 'xls', 'pdf', 'docx', 'doc']
    
    def process_whatsapp_document(self, media_id: str, user_id: str) -> Dict:
        """Process document uploaded via WhatsApp with enhanced capabilities"""
        try:
            # Download file from WhatsApp
            file_data = self._download_whatsapp_media(media_id)
            
            if not file_data:
                return {"status": "error", "message": "Could not download file"}
            
            # Determine file type and process
            filename = file_data.get('filename', '').lower()
            mime_type = file_data.get('mime_type', '').lower()
            file_size = file_data.get('size', 0)
            
            logging.info(f"Processing file: {filename} ({file_size} bytes, {mime_type})")
            
            # File size check (max 10MB)
            if file_size > 10 * 1024 * 1024:
                return {
                    "status": "error",
                    "message": "File too large (max 10MB). Please upload a smaller file."
                }
            
            # Route to appropriate processor
            if filename.endswith('.csv') or 'csv' in mime_type:
                return self._process_csv_enhanced(file_data['content'], user_id, filename)
            
            elif filename.endswith(('.xlsx', '.xls')) or 'excel' in mime_type or 'spreadsheet' in mime_type:
                return self._process_excel_enhanced(file_data['content'], user_id, filename)
            
            elif filename.endswith('.pdf') or 'pdf' in mime_type:
                return self._process_pdf_document(file_data['content'], user_id, filename)
            
            elif filename.endswith(('.docx', '.doc')) or 'word' in mime_type or 'document' in mime_type:
                return self._process_docx_document(file_data['content'], user_id, filename)
            
            else:
                return {
                    "status": "error",
                    "message": f"Unsupported file type. Supported: CSV, Excel, PDF, Word documents."
                }
                
        except Exception as e:
            logging.error(f"Enhanced file processing error: {e}")
            return {"status": "error", "message": f"Error processing file: {str(e)}"}
    
    def _download_whatsapp_media(self, media_id: str) -> Optional[Dict]:
        """Download media file from WhatsApp"""
        try:
            headers = {
                "Authorization": f"Bearer {current_app.config['ACCESS_TOKEN']}"
            }
            
            # Get media URL
            url = f"https://graph.facebook.com/{current_app.config['VERSION']}/{media_id}"
            response = requests.get(url, headers=headers, timeout=30)
            
            if response.status_code != 200:
                logging.error(f"Failed to get media info: {response.status_code}")
                return None
            
            media_info = response.json()
            media_url = media_info.get('url')
            
            if not media_url:
                logging.error("No media URL in response")
                return None
            
            # Download actual file
            file_response = requests.get(media_url, headers=headers, timeout=60)
            
            if file_response.status_code != 200:
                logging.error(f"Failed to download file: {file_response.status_code}")
                return None
            
            return {
                'content': file_response.content,
                'mime_type': media_info.get('mime_type', ''),
                'filename': media_info.get('filename', 'uploaded_file'),
                'size': len(file_response.content)
            }
            
        except Exception as e:
            logging.error(f"Error downloading WhatsApp media: {e}")
            return None
    
    def _process_pdf_document(self, file_content: bytes, user_id: str, filename: str) -> Dict:
        """Process PDF document and extract data"""
        try:
            extracted_data = []
            tables_found = []
            text_content = ""
            
            # Method 1: Try pdfplumber for tables
            try:
                with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                    for page_num, page in enumerate(pdf.pages):
                        # Extract text
                        page_text = page.extract_text()
                        if page_text:
                            text_content += f"\n--- Page {page_num + 1} ---\n{page_text}"
                        
                        # Extract tables
                        tables = page.extract_tables()
                        for table_num, table in enumerate(tables):
                            if table and len(table) > 1:  # At least header + one row
                                df = pd.DataFrame(table[1:], columns=table[0])
                                tables_found.append({
                                    'page': page_num + 1,
                                    'table': table_num + 1,
                                    'dataframe': df
                                })
                
                logging.info(f"Extracted {len(tables_found)} tables from PDF")
                
            except Exception as e:
                logging.warning(f"pdfplumber extraction failed: {e}")
            
            # Method 2: Fallback to PyPDF2 for text
            if not text_content:
                try:
                    pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                    for page_num, page in enumerate(pdf_reader.pages):
                        page_text = page.extract_text()
                        text_content += f"\n--- Page {page_num + 1} ---\n{page_text}"
                except Exception as e:
                    logging.warning(f"PyPDF2 extraction failed: {e}")
            
            # Process extracted tables
            if tables_found:
                best_table = self._find_best_data_table(tables_found)
                if best_table:
                    return self._process_dataframe(best_table['dataframe'], user_id, filename, 'pdf_table')
            
            # Process text content for sales data
            if text_content:
                sales_data = self._extract_sales_from_text(text_content)
                if sales_data:
                    return self._process_extracted_sales_data(sales_data, user_id, filename, 'pdf_text')
            
            # If no structured data found, return text analysis
            return {
                "status": "warning",
                "message": f"Could not find structured sales data in PDF. Extracted {len(text_content)} characters of text.",
                "data": {
                    "filename": filename,
                    "type": "pdf_text",
                    "text_length": len(text_content),
                    "tables_found": len(tables_found),
                    "suggestions": [
                        "PDF contains mostly text - try uploading a CSV/Excel file",
                        "If sales data is in the text, please copy/paste key information",
                        "Consider converting PDF tables to CSV format"
                    ]
                }
            }
            
        except Exception as e:
            logging.error(f"PDF processing error: {e}")
            return {"status": "error", "message": f"Error processing PDF: {str(e)}"}
    
    def _process_docx_document(self, file_content: bytes, user_id: str, filename: str) -> Dict:
        """Process DOCX document and extract data"""
        try:
            doc = Document(io.BytesIO(file_content))
            
            # Extract text from paragraphs
            text_content = ""
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Extract data from tables
            tables_data = []
            for table_num, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                
                if len(table_data) > 1:  # At least header + one row
                    try:
                        df = pd.DataFrame(table_data[1:], columns=table_data[0])
                        tables_data.append({
                            'table_num': table_num + 1,
                            'dataframe': df
                        })
                    except Exception as e:
                        logging.warning(f"Could not convert table {table_num + 1} to DataFrame: {e}")
            
            logging.info(f"Extracted {len(tables_data)} tables from DOCX")
            
            # Process tables first
            if tables_data:
                best_table = self._find_best_data_table(tables_data)
                if best_table:
                    return self._process_dataframe(best_table['dataframe'], user_id, filename, 'docx_table')
            
            # Process text content for sales data
            if text_content.strip():
                sales_data = self._extract_sales_from_text(text_content)
                if sales_data:
                    return self._process_extracted_sales_data(sales_data, user_id, filename, 'docx_text')
            
            # If no structured data found
            return {
                "status": "warning",
                "message": f"Could not find structured sales data in Word document.",
                "data": {
                    "filename": filename,
                    "type": "docx_text",
                    "text_length": len(text_content),
                    "tables_found": len(tables_data),
                    "suggestions": [
                        "Word document processed but no clear sales data found",
                        "Try uploading data in CSV or Excel format",
                        "Ensure sales data is in table format within the document"
                    ]
                }
            }
            
        except Exception as e:
            logging.error(f"DOCX processing error: {e}")
            return {"status": "error", "message": f"Error processing Word document: {str(e)}"}
    
    def _extract_sales_from_text(self, text: str) -> List[Dict]:
        """Extract sales data from plain text using regex patterns"""
        try:
            sales_data = []
            lines = text.split('\n')
            
            # Patterns to match sales data
            patterns = [
                # Date Product Quantity Price Total
                r'(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})\s+([^$\d]+?)\s+(\d+(?:\.\d+)?)\s+\$?(\d+(?:\.\d+)?)\s+\$?(\d+(?:\.\d+)?)',
                # Product: $Amount on Date
                r'([^:$\d]+?):\s*\$?(\d+(?:\.\d+)?)\s+(?:on|dated?)\s+(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})',
                # Sold X Product for $Y
                r'[Ss]old\s+(\d+(?:\.\d+)?)\s+([^$\d]+?)\s+(?:for|@)\s*\$?(\d+(?:\.\d+)?)',
            ]
            
            for line in lines:
                line = line.strip()
                if not line or len(line) < 10:
                    continue
                
                for pattern in patterns:
                    matches = re.findall(pattern, line, re.IGNORECASE)
                    for match in matches:
                        try:
                            if len(match) == 5:  # Full pattern
                                date_str, product, qty, price, total = match
                                sales_record = {
                                    'date': pd.to_datetime(date_str).to_pydatetime(),
                                    'product_name': product.strip(),
                                    'quantity': float(qty),
                                    'unit_price': float(price),
                                    'total_amount': float(total),
                                    'source': 'text_extraction'
                                }
                            elif len(match) == 3:  # Simplified patterns
                                if re.search(r'\d{1,2}[/-]\d{1,2}', match[2]):  # Date is third
                                    product, amount, date_str = match
                                    sales_record = {
                                        'date': pd.to_datetime(date_str).to_pydatetime(),
                                        'product_name': product.strip(),
                                        'quantity': 1.0,
                                        'unit_price': float(amount),
                                        'total_amount': float(amount),
                                        'source': 'text_extraction'
                                    }
                                else:  # Quantity, product, amount
                                    qty, product, amount = match
                                    sales_record = {
                                        'date': datetime.utcnow(),
                                        'product_name': product.strip(),
                                        'quantity': float(qty),
                                        'unit_price': float(amount) / float(qty),
                                        'total_amount': float(amount),
                                        'source': 'text_extraction'
                                    }
                            
                            # Validate record
                            if (sales_record['total_amount'] > 0 and 
                                sales_record['product_name'] and 
                                len(sales_record['product_name']) > 1):
                                sales_data.append(sales_record)
                                
                        except (ValueError, ZeroDivisionError) as e:
                            logging.debug(f"Failed to parse match {match}: {e}")
                            continue
            
            logging.info(f"Extracted {len(sales_data)} sales records from text")
            return sales_data
            
        except Exception as e:
            logging.error(f"Text extraction error: {e}")
            return []
    
    def _find_best_data_table(self, tables: List[Dict]) -> Optional[Dict]:
        """Find the table most likely to contain sales data"""
        try:
            best_table = None
            best_score = 0
            
            for table_info in tables:
                df = table_info['dataframe']
                score = 0
                
                # Check for sales-related columns
                column_str = ' '.join(df.columns.astype(str).str.lower())
                sales_keywords = ['date', 'product', 'quantity', 'price', 'amount', 'total', 'sale', 'revenue']
                
                for keyword in sales_keywords:
                    if keyword in column_str:
                        score += 1
                
                # Prefer tables with more rows
                score += min(len(df), 10) * 0.1
                
                # Prefer tables with numeric columns
                numeric_cols = df.select_dtypes(include=['number']).columns
                score += len(numeric_cols) * 0.5
                
                if score > best_score:
                    best_score = score
                    best_table = table_info
            
            return best_table if best_score > 1 else None
            
        except Exception as e:
            logging.error(f"Error finding best table: {e}")
            return None
    
    def _process_csv_enhanced(self, file_content: bytes, user_id: str, filename: str) -> Dict:
        """Enhanced CSV processing with better error handling"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            df = None
            
            for encoding in encodings:
                try:
                    csv_string = file_content.decode(encoding)
                    df = pd.read_csv(io.StringIO(csv_string))
                    logging.info(f"Successfully read CSV with {encoding} encoding")
                    break
                except UnicodeDecodeError:
                    continue
                except Exception as e:
                    logging.warning(f"Failed to read CSV with {encoding}: {e}")
                    continue
            
            if df is None:
                return {"status": "error", "message": "Could not read CSV file with any supported encoding"}
            
            return self._process_dataframe(df, user_id, filename, 'csv')
            
        except Exception as e:
            logging.error(f"Enhanced CSV processing error: {e}")
            return {"status": "error", "message": f"Error reading CSV: {str(e)}"}
    
    def _process_excel_enhanced(self, file_content: bytes, user_id: str, filename: str) -> Dict:
        """Enhanced Excel processing with multiple sheet support"""
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(io.BytesIO(file_content))
            sheet_names = excel_file.sheet_names
            
            logging.info(f"Excel file has {len(sheet_names)} sheets: {sheet_names}")
            
            # Try to find the best sheet with sales data
            best_sheet = None
            best_score = 0
            
            for sheet_name in sheet_names:
                try:
                    df = pd.read_excel(io.BytesIO(file_content), sheet_name=sheet_name)
                    
                    # Score sheet based on sales-related content
                    score = 0
                    column_str = ' '.join(df.columns.astype(str).str.lower())
                    sales_keywords = ['date', 'product', 'quantity', 'price', 'amount', 'total', 'sale']
                    
                    for keyword in sales_keywords:
                        if keyword in column_str:
                            score += 1
                    
                    # Prefer sheets with more data
                    score += min(len(df), 50) * 0.01
                    
                    if score > best_score:
                        best_score = score
                        best_sheet = df
                        
                except Exception as e:
                    logging.warning(f"Could not read sheet {sheet_name}: {e}")
                    continue
            
            if best_sheet is not None:
                return self._process_dataframe(best_sheet, user_id, filename, 'excel')
            else:
                # Fallback to first sheet
                df = pd.read_excel(io.BytesIO(file_content))
                return self._process_dataframe(df, user_id, filename, 'excel')
            
        except Exception as e:
            logging.error(f"Enhanced Excel processing error: {e}")
            return {"status": "error", "message": f"Error reading Excel file: {str(e)}"}
    
    def _process_extracted_sales_data(self, sales_data: List[Dict], user_id: str, filename: str, source_type: str) -> Dict:
        """Process sales data extracted from text"""
        try:
            from .sales_models import SalesDataManager
            sales_manager = SalesDataManager(self.db_manager)
            
            if not sales_data:
                return {
                    "status": "warning",
                    "message": "No sales data could be extracted from the document",
                    "data": {"filename": filename, "type": source_type}
                }
            
            # Save sales records
            result = sales_manager.save_bulk_sales_data(user_id, sales_data)
            
            # Track file upload event
            if hasattr(self.db_manager, 'track_event'):
                self.db_manager.track_event("file_upload", user_id, {
                    "filename": filename,
                    "type": source_type,
                    "records_processed": len(sales_data),
                    "success_count": result.get("success_count", 0),
                    "error_count": result.get("error_count", 0)
                })
            
            return {
                "status": "success",
                "message": f"Extracted and processed {result['success_count']} sales records from {filename}",
                "data": {
                    "filename": filename,
                    "type": source_type,
                    "total_processed": result["total_processed"],
                    "success_count": result["success_count"],
                    "error_count": result["error_count"],
                    "errors": result["errors"][:5] if result["errors"] else []
                }
            }
            
        except Exception as e:
            logging.error(f"Error processing extracted sales data: {e}")
            return {"status": "error", "message": f"Error processing extracted data: {str(e)}"}
    
    # Keep existing _process_dataframe method from previous implementation...
    def _process_dataframe(self, df: pd.DataFrame, user_id: str, filename: str, file_type: str) -> Dict:
        """Process pandas DataFrame and extract sales data"""
        try:
            # Clean column names
            df.columns = df.columns.str.strip().str.lower()
            
            # Log DataFrame info
            logging.info(f"Processing {file_type} file: {filename}")
            logging.info(f"Columns: {list(df.columns)}")
            logging.info(f"Rows: {len(df)}")
            
            # Detect data type based on columns
            if self._is_sales_data(df.columns):
                return self._process_sales_data(df, user_id, filename)
            elif self._is_product_data(df.columns):
                return self._process_product_data(df, user_id, filename)
            else:
                return self._attempt_generic_processing(df, user_id, filename)
                
        except Exception as e:
            logging.error(f"DataFrame processing error: {e}")
            return {"status": "error", "message": f"Error processing data: {str(e)}"}
    
    def _is_sales_data(self, columns: List[str]) -> bool:
        """Check if DataFrame contains sales data"""
        sales_indicators = [
            'date', 'product', 'quantity', 'price', 'amount', 'total', 
            'customer', 'sale', 'revenue', 'item'
        ]
        
        column_str = ' '.join(columns)
        matches = sum(1 for indicator in sales_indicators if indicator in column_str)
        return matches >= 2
    
    def _is_product_data(self, columns: List[str]) -> bool:
        """Check if DataFrame contains product data"""
        product_indicators = [
            'product', 'name', 'price', 'cost', 'stock', 'inventory', 
            'category', 'sku', 'description'
        ]
        
        column_str = ' '.join(columns)
        matches = sum(1 for indicator in product_indicators if indicator in column_str)
        return matches >= 2
    
    def _process_sales_data(self, df: pd.DataFrame, user_id: str, filename: str) -> Dict:
        """Process sales data from DataFrame"""
        try:
            from .sales_models import SalesDataManager
            sales_manager = SalesDataManager(self.db_manager)
            
            # Map columns to expected fields
            column_mapping = self._create_sales_column_mapping(df.columns)
            
            sales_records = []
            errors = []
            
            for index, row in df.iterrows():
                try:
                    sales_record = self._extract_sales_record(row, column_mapping)
                    if sales_record:
                        sales_records.append(sales_record)
                    else:
                        errors.append(f"Row {index + 1}: Could not extract valid sales data")
                        
                except Exception as e:
                    errors.append(f"Row {index + 1}: {str(e)}")
            
            if not sales_records:
                return {
                    "status": "error",
                    "message": "No valid sales records found in file",
                    "errors": errors[:10]  # Show first 10 errors
                }
            
            # Save sales records
            result = sales_manager.save_bulk_sales_data(user_id, sales_records)
            
            # Track file upload event
            if hasattr(self.db_manager, 'track_event'):
                self.db_manager.track_event("file_upload", user_id, {
                    "filename": filename,
                    "type": "sales_data",
                    "records_processed": len(sales_records),
                    "success_count": result.get("success_count", 0),
                    "error_count": result.get("error_count", 0)
                })
            
            return {
                "status": "success",
                "message": f"Processed {result['success_count']} sales records from {filename}",
                "data": {
                    "filename": filename,
                    "type": "sales_data",
                    "total_processed": result["total_processed"],
                    "success_count": result["success_count"],
                    "error_count": result["error_count"],
                    "errors": result["errors"][:5] if result["errors"] else []
                }
            }
            
        except Exception as e:
            logging.error(f"Sales data processing error: {e}")
            return {"status": "error", "message": f"Error processing sales data: {str(e)}"}
    
    def _create_sales_column_mapping(self, columns: List[str]) -> Dict:
        """Create mapping from DataFrame columns to sales record fields"""
        mapping = {}
        
        # Common column name variations
        date_cols = ['date', 'sale_date', 'transaction_date', 'order_date']
        product_cols = ['product', 'product_name', 'item', 'item_name', 'name']
        quantity_cols = ['quantity', 'qty', 'amount', 'units']
        price_cols = ['price', 'unit_price', 'cost', 'rate']
        total_cols = ['total', 'total_amount', 'revenue', 'sales', 'value']
        customer_cols = ['customer', 'customer_name', 'client', 'buyer']
        
        for col in columns:
            if any(date_col in col for date_col in date_cols):
                mapping['date'] = col
            elif any(product_col in col for product_col in product_cols):
                mapping['product_name'] = col
            elif any(qty_col in col for qty_col in quantity_cols):
                mapping['quantity'] = col
            elif any(price_col in col for price_col in price_cols):
                mapping['unit_price'] = col
            elif any(total_col in col for total_col in total_cols):
                mapping['total_amount'] = col
            elif any(customer_col in col for customer_col in customer_cols):
                mapping['customer_name'] = col
        
        return mapping
    
    def _extract_sales_record(self, row: pd.Series, column_mapping: Dict) -> Optional[Dict]:
        """Extract sales record from DataFrame row"""
        try:
            record = {}
            
            # Extract date
            if 'date' in column_mapping:
                date_value = row[column_mapping['date']]
                if pd.notna(date_value):
                    if isinstance(date_value, str):
                        record['date'] = pd.to_datetime(date_value).to_pydatetime()
                    else:
                        record['date'] = date_value
                else:
                    record['date'] = datetime.utcnow()
            else:
                record['date'] = datetime.utcnow()
            
            # Extract product name
            if 'product_name' in column_mapping:
                product = row[column_mapping['product_name']]
                record['product_name'] = str(product) if pd.notna(product) else ""
            else:
                record['product_name'] = "Unknown Product"
            
            # Extract quantity
            if 'quantity' in column_mapping:
                qty = row[column_mapping['quantity']]
                record['quantity'] = float(qty) if pd.notna(qty) and qty != "" else 1.0
            else:
                record['quantity'] = 1.0
            
            # Extract unit price
            if 'unit_price' in column_mapping:
                price = row[column_mapping['unit_price']]
                record['unit_price'] = float(price) if pd.notna(price) and price != "" else 0.0
            else:
                record['unit_price'] = 0.0
            
            # Extract total amount
            if 'total_amount' in column_mapping:
                total = row[column_mapping['total_amount']]
                record['total_amount'] = float(total) if pd.notna(total) and total != "" else 0.0
            else:
                record['total_amount'] = record['quantity'] * record['unit_price']
            
            # Extract customer name
            if 'customer_name' in column_mapping:
                customer = row[column_mapping['customer_name']]
                record['customer_name'] = str(customer) if pd.notna(customer) else ""
            else:
                record['customer_name'] = ""
            
            # Additional fields
            record['source'] = 'file_upload'
            record['category'] = 'general'
            record['payment_method'] = 'unknown'
            record['notes'] = f"Imported from file"
            
            # Validate record
            if record['total_amount'] <= 0 and record['unit_price'] <= 0:
                return None
            
            return record
            
        except Exception as e:
            logging.error(f"Error extracting sales record: {e}")
            return None
    
    def _process_product_data(self, df: pd.DataFrame, user_id: str, filename: str) -> Dict:
        """Process product data from DataFrame"""
        # Implementation remains the same as before
        return {
            "status": "success",
            "message": f"Product data processing not yet fully implemented for {filename}",
            "data": {"filename": filename, "type": "product_data"}
        }
    
    def _attempt_generic_processing(self, df: pd.DataFrame, user_id: str, filename: str) -> Dict:
        """Attempt to process unknown data format"""
        # Implementation remains the same as before
        try:
            suggestions = []
            numeric_cols = df.select_dtypes(include=['number']).columns.tolist()
            date_cols = []
            text_cols = df.select_dtypes(include=['object']).columns.tolist()
            
            # Try to find date columns
            for col in text_cols:
                try:
                    pd.to_datetime(df[col].dropna().iloc[0])
                    date_cols.append(col)
                except:
                    pass
            
            if len(numeric_cols) >= 2 and len(text_cols) >= 1:
                suggestions.append("This might be sales data. Expected columns: date, product_name, quantity, unit_price, total_amount")
            
            return {
                "status": "warning",
                "message": f"Could not automatically detect data type in {filename}",
                "data": {
                    "filename": filename,
                    "columns": list(df.columns),
                    "numeric_columns": numeric_cols,
                    "date_columns": date_cols,
                    "text_columns": text_cols,
                    "row_count": len(df),
                    "suggestions": suggestions
                }
            }
            
        except Exception as e:
            logging.error(f"Generic processing error: {e}")
            return {"status": "error", "message": f"Error analyzing file: {str(e)}"}


# Initialize enhanced file processor
def create_enhanced_file_processor(db_manager):
    return EnhancedFileProcessor(db_manager)